import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_callout_box(doc, text_paragraphs, title="KEY ARCHITECTURAL HIGHLIGHT"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F1F5F9")
    set_cell_margins(cell, top=160, bottom=160, left=200, right=200)
    
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>\n'
        f'  <w:left w:val="single" w:sz="36" w:space="0" w:color="0F766E"/>\n'
        f'  <w:top w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'  <w:bottom w:val="none"/>\n'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_before = Pt(2)
    p0.paragraph_format.space_after = Pt(4)
    run_t = p0.add_run(f"💡 {title}")
    run_t.font.name = 'Calibri'
    run_t.font.size = Pt(11)
    run_t.font.bold = True
    run_t.font.color.rgb = RGBColor(15, 118, 110)
    
    for text in text_paragraphs:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(51, 65, 85)

def build_overview_document(output_path):
    doc = docx.Document()
    
    # Page setup - Standard 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)

    # Color Palette
    PRIMARY = RGBColor(15, 118, 110)     # Deep Teal #0F766E
    SECONDARY = RGBColor(15, 23, 42)    # Slate Navy #0F172A
    BODY_TEXT = RGBColor(51, 65, 85)    # Slate 700 #334155
    MUTED = RGBColor(100, 116, 139)     # Slate 500 #64748B
    DARK_BLUE = RGBColor(30, 58, 138)   # Deep Blue #1E3A8A

    # Set default style font
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = BODY_TEXT

    def add_doc_header():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        r_meta = p.add_run("MEDIUNITY — CLOUD-READY DISTRIBUTED HEALTHCARE & WELLNESS PLATFORM")
        r_meta.font.size = Pt(9.5)
        r_meta.font.bold = True
        r_meta.font.color.rgb = MUTED
        
        p_main = doc.add_paragraph()
        p_main.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_main.paragraph_format.space_before = Pt(4)
        p_main.paragraph_format.space_after = Pt(6)
        r_title = p_main.add_run("PROJECT OVERVIEW AND DISCUSSION")
        r_title.font.size = Pt(24)
        r_title.font.bold = True
        r_title.font.color.rgb = PRIMARY
        
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_before = Pt(0)
        p_sub.paragraph_format.space_after = Pt(20)
        r_sub = p_sub.add_run("Exhaustive Core Functional Modules & User Workflows Breakdown (Part-by-Part)")
        r_sub.font.size = Pt(12)
        r_sub.font.italic = True
        r_sub.font.color.rgb = SECONDARY

        # Horizontal Divider line
        p_div = doc.add_paragraph()
        p_div.paragraph_format.space_after = Pt(18)
        p_div_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="1" w:color="0F766E"/></w:pBdr>')
        p_div._element.get_or_add_pPr().append(p_div_border)

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(22)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.size = Pt(17)
        run.font.bold = True
        run.font.color.rgb = PRIMARY
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.size = Pt(13.5)
        run.font.bold = True
        run.font.color.rgb = SECONDARY
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = DARK_BLUE
        return p

    def add_p(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = BODY_TEXT
        return p

    def add_bullet(lead, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run_lead = p.add_run(lead + " ")
        run_lead.font.bold = True
        run_lead.font.color.rgb = SECONDARY
        run_text = p.add_run(text)
        run_text.font.color.rgb = BODY_TEXT
        return p

    # --- DOCUMENT CONTENT GENERATION ---
    add_doc_header()

    # SECTION 1
    add_h1("1. Executive Project Summary & Mission Statement")
    add_p("MediUnity (also referenced as the Medicare Systems Platform) is an enterprise-grade, cloud-native, multi-portal digital healthcare and wellness ecosystem. Built with a scalable multi-tenant architecture, MediUnity seamlessly connects healthcare consumers (patients), clinical professionals (medical doctors), healthcare facilities (hospitals and diagnostic centers), and pharmaceutical providers (retail pharmacies) into a single, unified digital platform.")
    add_p("The core mission of MediUnity is to eliminate traditional healthcare operational silos and geographical accessibility barriers by engineering a globally accessible platform that digitizes the entire spectrum of outpatient medical care—from preliminary doctor discovery and instant appointment booking to real-time telehealth consultations, electronic health record (EHR) management, lab test ordering, and pharmacy prescription fulfillment.")
    
    add_callout_box(doc, [
        "Global Scope & Scalability: MediUnity is architected for worldwide deployment across diverse healthcare ecosystems. It avoids localized hardcoding, supporting international metric formats, multi-currency processing via Stripe, and distributed multi-portal rollouts suitable for international clinic networks and hospitals.",
        "Unified Multi-Tenant Core: Rather than forcing diverse stakeholders into a generic monolithic interface, MediUnity deploys specialized user-role applications (Patient Portal, Doctor Portal, Partner Portal, Admin Control Center) backed by a central, high-throughput Node.js orchestration engine."
    ], title="EXECUTIVE ARCHITECTURAL PRINCIPLE")

    # SECTION 2
    add_h1("2. Background, Context & Problem Statement")
    add_p("Global healthcare delivery faces mounting operational challenges driven by fragmented digital tools, manual paper-based recordkeeping, physical clinic queue bottlenecks, and isolated service providers. Patients and healthcare institutions routinely suffer from systemic inefficiencies:")
    
    add_bullet("Fragmented Patient Records:", "Patient health histories, diagnostic laboratory reports, and prescription records remain locked inside disconnected clinical databases, preventing holistic patient evaluation and increasing medical diagnostic errors.")
    add_bullet("Scheduling Friction & Queue Delays:", "Manual clinic phone bookings and in-person queue management lead to prolonged waiting times, underutilized doctor schedules, and high appointment no-show rates.")
    add_bullet("Diagnostic Laboratory Isolation:", "Patients face cumbersome processes to find accredited diagnostic centers, schedule sample pickups, and receive digital lab results securely.")
    add_bullet("Pharmacy Inventory & Prescription Delays:", "Manual paper prescriptions lead to poor drug adherence, inventory stockout frustration, and delayed home delivery of essential chronic medications.")
    add_bullet("Lack of Unified Regulatory Oversight:", "System administrators lack real-time telemetry, credential verification audit trails, and automated financial settlement pipelines to govern multi-provider networks effectively.")

    add_p("MediUnity resolves these systemic failures by establishing a unified digital bridge across all healthcare participants, enforcing strict credential verification, zero-trust role-based access control (RBAC), and automated background service queues.")

    # SECTION 3 - CORE FUNCTIONAL MODULES & USER WORKFLOWS (DETAILED PART BY PART)
    add_h1("3. Comprehensive Core Functional Modules & User Workflows (Part-by-Part)")
    add_p("MediUnity organizes its clinical and operational features into five interconnected functional pillars, described in detail below:")

    # PART 1
    add_h2("Part 1: Patient Ecosystem & Telehealth Experience")
    add_p("The Patient Ecosystem (comprising frontend and frontend-patient applications) delivers an intuitive, accessible experience for healthcare consumers globally.")

    add_h3("1.1 Authentication, Role Gateway & Account Management")
    add_bullet("Multi-Role Registration & Login:", "Supports email/password authentication backed by bcrypt password hashing and JWT token issuance (Access Tokens + HTTP-only Refresh Tokens).")
    add_bullet("Portal Gateway Switcher:", "Allows authenticated users to seamlessly switch between Patient, Doctor, or Partner profiles depending on authorized account permissions.")
    add_bullet("Security Session Rotation:", "Integrates RefreshToken models for secure token invalidation and device session revocation.")

    add_h3("1.2 Patient Medical Profile & Family Dependent Vault")
    add_bullet("Personal EHR Demographics:", "Stores patient personal metrics (Blood Group, Height, Weight, Emergency Contacts, Allergic History, Chronic Illnesses, Insurance Policy numbers).")
    add_bullet("Family Member Profiles:", "Allows primary patient account holders to create secondary dependent profiles (e.g., children, elderly parents) to schedule appointments and manage prescriptions on their behalf.")

    add_h3("1.3 Smart Doctor Discovery & Consultation Booking Workflow")
    add_bullet("Multi-Criteria Search & Filter Engine:", "Enables patients to query doctors by medical specialty (e.g., Cardiology, Neurology, Pediatrics, Dermatology), location, consultation fee range, user ratings, and gender.")
    add_bullet("Doctor Profile & Real-Time Availability:", "Displays doctor bio, qualifications, hospital affiliations, consultation fees, user reviews, and interactive calendar availability slots.")
    add_bullet("Real-Time Slot Locking & Stripe Checkout:", "When a patient selects a time slot, the system locks the slot to prevent double-booking. Payment is processed securely via Stripe API. Upon payment authorization, the appointment transitions to confirmed status.")
    add_bullet("Multi-Channel Alert Dispatches:", "Triggers asynchronous notification workers via BullMQ, sending instant browser dispatches, Nodemailer email receipts, and Twilio SMS reminders.")

    add_h3("1.4 Interactive Teleconsultation & Real-Time Messaging Hub")
    add_bullet("Socket.IO WebSocket Telehealth Chat:", "Establishes low-latency, bi-directional real-time communication channels between doctor and patient during scheduled appointment slots.")
    add_bullet("Clinical File & Media Sharing:", "Allows patients to attach past diagnostic test images, lab reports, or symptom photographs during live chat using Cloudinary secure media pipelines.")
    add_bullet("Live Doctor Presence & Queue Status:", "Displays live doctor online/offline indicators and estimated waiting queue duration.")

    add_h3("1.5 Electronic Health Record (EHR) & Vitals Health Tracker")
    add_bullet("Continuous Vitals Logging:", "Patients log daily health metrics: Blood Pressure (Systolic/Diastolic), Blood Glucose (Fasting/Post-prandial), Heart Rate (BPM), Body Temperature (°F/°C), and BMI.")
    add_bullet("Interactive Trend Visualizations:", "Render responsive time-series graphs displaying historical vitals trends to identify health spikes or chronic degradation.")
    add_bullet("EHR Document Vault:", "Stores downloadable PDF prescriptions issued by doctors, diagnostic laboratory test results, and uploaded medical file archives.")

    add_h3("1.6 AI Symptom Checker & Health Journal")
    add_bullet("Interactive Symptom Triage Tool:", "Guides patients through a structured clinical symptom questionnaire, analyzing reported discomforts and recommending suitable doctor specialties.")
    add_bullet("Personal Mood & Wellness Journal:", "Private health diary allowing patients to record daily mood logs, stress levels, mental wellness notes, and physical exercise routines.")

    add_h3("1.7 Pharmacy E-Commerce & Medication Delivery Workflow")
    add_bullet("Medicine Search & Inventory Catalog:", "Browse prescription medications, OTC drugs, health supplements, and healthcare equipment with real-time pricing.")
    add_bullet("OCR Prescription Parsing (Tesseract.js):", "Patients upload paper prescription images, which are automatically parsed via Tesseract.js optical character recognition to extract medication names.")
    add_bullet("Pharmacy Selection & Order Checkout:", "Patients choose local or online partner pharmacies, confirm delivery addresses, upload verified prescriptions, and complete online payment.")
    add_bullet("Order Fulfillment & Delivery Tracking:", "Tracks order status through real-time state transitions (Order Placed -> Verified -> Packed -> Out for Delivery -> Delivered).")

    add_h3("1.8 Diagnostic Center & Lab Test Booking Workflow")
    add_bullet("Diagnostic Test Cataloging:", "Explore comprehensive diagnostic test menus (Complete Blood Count, Lipid Profile, MRI Scans, Ultrasound, Thyroid Panel).")
    add_bullet("Sample Collection Scheduling:", "Choose between home specimen pickup or diagnostic center walk-in appointment slots.")
    add_bullet("Electronic Result Delivery:", "Diagnostic centers upload completed lab test PDF reports directly to the patient's EHR vault, sending instant push and SMS notifications.")

    add_h3("1.9 Home Healthcare Services & Emergency Ambulance Tracking")
    add_bullet("Home Nursing & Healthcare Booking:", "Schedule home visits for post-operative nursing care, physiotherapy, elderly care assistance, and wound dressing.")
    add_bullet("Real-Time Ambulance Dispatch Tracking:", "Requests emergency medical transport and tracks real-time GPS location updates of assigned emergency vehicles.")

    # PART 2
    add_h2("Part 2: Doctor Clinical Workspace & Practice Management")
    add_p("The Doctor Portal (frontend-doctor) provides medical clinicians with comprehensive tools to manage consultations, issue prescriptions, review medical histories, and track clinical earnings.")

    add_h3("2.1 Clinician Onboarding & Credential Verification")
    add_bullet("License & Qualification Submission:", "Prospective doctors submit medical registration numbers (BMDC / Medical Council), specialization degrees, institutional affiliations, and government identity proofs.")
    add_bullet("Verification Status Governance:", "Accounts remain in unverified state until platform administrators audit submitted credentials, granting isVerified status to publish profiles publicly.")

    add_h3("2.2 Dynamic Availability Calendar & Slot Builder")
    add_bullet("Recurring Consultation Slots:", "Doctors define weekly recurring work shifts, consultation slot durations (15, 20, 30 mins), consultation fees, and emergency buffers.")
    add_bullet("Slot Blockade & Schedule Adjustments:", "Enables doctors to temporarily mark unavailability or cancel slots, automatically triggering patient refunds and reschedule dispatches.")

    add_h3("2.3 Clinical Consultation Roster & Patient EHR Inspection")
    add_bullet("Consultation Roster Dashboard:", "Chronological schedule of upcoming, completed, and canceled patient appointments with one-click access to consultation rooms.")
    add_bullet("Pre-Consultation Medical History Review:", "Doctors inspect patient vitals trend graphs, past prescriptions, allergic histories, and diagnostic lab reports prior to initiating telehealth sessions.")

    add_h3("2.4 Digital Prescription Builder (PrescriptionBuilder)")
    add_bullet("Structured Clinical Scripting:", "Doctors input primary diagnosis, ICD disease codes, drug dosage instructions (1-0-1, before/after meals), duration, diet advice, and recommended lab tests.")
    add_bullet("PDF Generation & Instant EHR Vault Sync:", "Auto-generates digital PDF prescriptions, stores them in MongoDB, and instantly syncs them to the patient's mobile app and pharmacy ordering queue.")

    add_h3("2.5 Specialist Referral & Hospital Transfer System")
    add_bullet("Inter-Specialist Referral Notes:", "Refers patients to sub-specialists or tertiary hospitals with attached clinical summary notes and diagnostic findings.")

    add_h3("2.6 Financial Accounting & Doctor Payouts")
    add_bullet("Earnings & Commission Accounting:", "Displays gross consultation revenue, platform commission deductions, and net withdrawal balances.")
    add_bullet("Payout Disbursement Requests:", "Submits payout withdrawal requests to platform administrators, tracking bank/mobile money transfer updates.")

    # PART 3
    add_h2("Part 3: Partner Provider Ecosystem (Hospitals, Diagnostic Centers, Pharmacies)")
    add_p("The Partner Portal (frontend-partner) equips institutional healthcare providers with operational management interfaces.")

    add_h3("3.1 Hospital Operations & Facility Management")
    add_bullet("Department & Staff Administration:", "Configure hospital clinical departments (Emergency, ICU, Cardiology, Pediatrics) and assign resident doctor rosters.")
    add_bullet("Inpatient Bed & ICU Emergency Inquiries:", "Manages real-time inquiries for available ICU beds, general ward rooms, and emergency surgery admissions.")

    add_h3("3.2 Diagnostic Laboratory Operations")
    add_bullet("Test Menu & Pricing Catalog:", "Publish available laboratory tests, specimen requirements, fasting guidelines, turnaround times, and pricing.")
    add_bullet("Specimen Dispatch & Digital Results Upload:", "Assign phlebotomists for home sample collection and upload digital lab test PDF reports directly to patient EHR accounts.")

    add_h3("3.3 Retail Pharmacy Management")
    add_bullet("Medicine Stock Inventory:", "Manage drug inventory quantities, unit prices, batch numbers, and stock expiration alerts.")
    add_bullet("Order Processing & Prescription Verification:", "Review incoming patient orders, verify doctor digital signatures, pack medications, and assign delivery riders.")

    # PART 4
    add_h2("Part 4: Administrative Governance, Security & System Telemetry")
    add_p("The Administrative Control Center (admin) provides platform moderators and system engineers with global oversight.")

    add_h3("4.1 Partner Verification & License Audit Desk")
    add_bullet("Credential Auditing Desk:", "Centralized verification dashboard to review medical license credentials, institutional registration certificates, and ID proofs for pending doctors, hospitals, labs, and pharmacies.")
    add_bullet("Approval / Rejection Governance:", "Approves verified partners (isVerified: true) to enable public marketplace listing or requests document resubmissions.")

    add_h3("4.2 Financial & Revenue Allocation Engine")
    add_bullet("Global Commission Control:", "Configures platform commission percentages across doctor fees, lab bookings, and pharmacy sales.")
    add_bullet("Payout Approvals Desk:", "Audits and approves financial payout disbursements to doctors and partner facilities.")

    add_h3("4.3 System Telemetry & Infrastructure Performance Monitoring")
    add_bullet("Real-Time Telemetry Gauges:", "Monitors server CPU/memory load, API endpoint latencies, database connection pool stats, and BullMQ background task throughput.")

    add_h3("4.4 Immutable Audit Logging")
    add_bullet("Security Audit Trails:", "Appends every administrative action, status modification, and role escalation into an immutable AuditLog collection for regulatory compliance auditing.")

    # PART 5
    add_h2("Part 5: Social Healthcare Community & Educational Content Hub")
    add_p("Promotes healthcare awareness and peer support across patients and medical professionals.")

    add_h3("5.1 Medical Community Forum & Peer Discussions")
    add_bullet("Health Discussion Boards:", "Patients and doctors share medical wellness advice, health tips, and community support posts.")
    add_bullet("Community Interaction Engine:", "Supports post liking, commenting, category filtering (Nutrition, Mental Health, Fitness), and content moderation reporting.")

    add_h3("5.2 Peer-Reviewed Health Articles & Wellness Blog")
    add_bullet("Expert Publishing Platform:", "Verified doctors write educational health articles, preventive medical advice, and wellness blogs.")

    # SECTION 4
    add_h1("4. Technical Discussion & Architectural Evaluation")
    add_p("This section evaluates the primary technical paradigms governing MediUnity:")
    add_bullet("Asynchronous Task Processing (Redis + BullMQ):", "Decouples heavy external dispatches (SMS via Twilio, Emails via Nodemailer, OCR via Tesseract.js) from the main Node.js event loop, guaranteeing sub-100ms API responses.")
    add_bullet("Defense-in-Depth Security Architecture:", "Combines JWT token authentication, Express RBAC middleware, NoSQL injection filtering via mongo-sanitize, Helmet security headers, and IP rate limiting.")
    add_bullet("Real-Time WebSockets Communication:", "Employs Socket.IO for low-latency doctor-patient telehealth chat, live queue status updates, and instant emergency notifications.")

    # SECTION 5
    add_h1("5. Comparative Feature Matrix")
    add_p("Evaluates MediUnity's technical advantages over legacy healthcare workflows:")

    comp_table = doc.add_table(rows=6, cols=4)
    comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    comp_table.autofit = False

    comp_headers = ["Feature Domain", "Traditional Paper Systems", "Single-Purpose Apps", "MediUnity Platform"]
    comp_widths = [Inches(1.5), Inches(1.6), Inches(1.7), Inches(1.7)]

    c_hdr_cells = comp_table.rows[0].cells
    for i, text in enumerate(comp_headers):
        c_hdr_cells[i].width = comp_widths[i]
        p = c_hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(c_hdr_cells[i], "0F766E")
        set_cell_margins(c_hdr_cells[i], top=120, bottom=120, left=120, right=120)

    comp_data = [
        ("Multi-Stakeholder Integration", "Manual paper handoffs, zero digital connection", "Doctor-Patient only; no pharmacy or lab integration", "Unified Patients, Doctors, Hospitals, Labs, Pharmacies & Admins"),
        ("Appointment Booking", "Physical queues & manual phone calls", "Basic calendar booking", "Real-time slot lock, Stripe payments, automated queue & alerts"),
        ("Health Record Storage", "Siloed physical paper charts", "Fragmented PDF uploads", "Centralized Cloud MongoDB EHR with vitals trend graphing"),
        ("Background Operations", "Manual staff execution", "Synchronous blocking HTTP loops", "Asynchronous Redis + BullMQ task queues for high throughput"),
        ("Security & Compliance", "Physical locks, high risk of record loss", "Basic HTTPS login", "Defense-in-depth: Helmet, NoSQL sanitization, RBAC & Audit Logs")
    ]

    for row_idx, data in enumerate(comp_data, start=1):
        row_cells = comp_table.rows[row_idx].cells
        bg_hex = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            row_cells[col_idx].width = comp_widths[col_idx]
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.font.size = Pt(9.5)
            if col_idx == 3:
                run.font.bold = True
                run.font.color.rgb = PRIMARY
            elif col_idx == 0:
                run.font.bold = True
                run.font.color.rgb = SECONDARY
            else:
                run.font.color.rgb = BODY_TEXT
            set_cell_background(row_cells[col_idx], bg_hex)
            set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=120, right=120)

    p_space2 = doc.add_paragraph()
    p_space2.paragraph_format.space_before = Pt(12)

    # SECTION 6
    add_h1("6. Implementation Challenges & Technical Mitigation Strategies")
    add_bullet("Challenge 1: Third-Party API Latencies.", "Offloaded Twilio SMS, Nodemailer Email, and Tesseract.js OCR to BullMQ worker queues.")
    add_bullet("Challenge 2: Multi-Portal State Sync.", "Implemented a central Socket.IO event router broadcasting room-based updates.")
    add_bullet("Challenge 3: Medical License Verification.", "Enforced dual-stage Admin verification workflow (isVerified: true) prior to public profile visibility.")
    add_bullet("Challenge 4: Diagnostic File Security.", "Configured Cloudinary private upload presets with signed temporal URLs.")

    # SECTION 7
    add_h1("7. Strategic Future Directions & Recommendations")
    add_bullet("AI Symptom Triage LLM Assistant:", "Automating specialty recommendations based on natural language symptom descriptions.")
    add_bullet("Direct Wearable IoT Sync:", "Ingesting continuous vitals (ECG, SpO2, Continuous Glucose) via Apple HealthKit and Google Fit APIs.")
    add_bullet("Embedded WebRTC Video Teleconsultations:", "Upgrading Socket.IO text chat to native peer-to-peer audio/video streaming.")

    # SECTION 8
    add_h1("8. Conclusion")
    add_p("MediUnity demonstrates a robust, scalable multi-portal digital healthcare architecture. By unifying Patients, Doctors, Hospitals, Labs, Pharmacies, and Admins under a single secure JavaScript stack, MediUnity delivers a comprehensive foundation for modern global digital health.")

    # Save Document
    doc.save(output_path)
    print(f"Successfully generated report document at: {output_path}")

if __name__ == "__main__":
    downloads_dir = os.path.join(os.path.expanduser("~"), "Downloads")
    os.makedirs(downloads_dir, exist_ok=True)
    target_file = os.path.join(downloads_dir, "MediUnity_Project_Overview_and_Discussion.docx")
    build_overview_document(target_file)
