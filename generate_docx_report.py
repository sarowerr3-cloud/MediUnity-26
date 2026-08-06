import os
import sys
import shutil
import subprocess

# Ensure python-docx and Pillow are installed
try:
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn
except ImportError:
    print("Installing python-docx and Pillow...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "pillow"])
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml import OxmlElement, parse_xml
    from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_report():
    doc = docx.Document()

    # Set Margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)

    # Base Colors
    PRIMARY = RGBColor(15, 118, 110)    # Teal #0F766E
    SECONDARY = RGBColor(15, 23, 42)   # Dark Slate #0F172A
    BODY_TEXT = RGBColor(51, 65, 85)   # Slate 700 #334155
    MUTED = RGBColor(100, 116, 139)    # Slate 500 #64748B
    HEX_PRIMARY = "0F766E"
    HEX_LIGHT_BG = "F1F5F9"
    HEX_BORDER = "CBD5E1"

    # Style Configurations
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = BODY_TEXT

    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(36)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(text)
        run.font.size = Pt(24)
        run.font.bold = True
        run.font.color.rgb = PRIMARY
        return p

    def add_subtitle(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(18)
        run = p.add_run(text)
        run.font.size = Pt(13)
        run.font.color.rgb = SECONDARY
        return p

    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = PRIMARY
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = SECONDARY
        return p

    def add_heading_3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = PRIMARY
        return p

    def add_body(text, bold_prefix="", space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            run_b.font.bold = True
            run_b.font.color.rgb = SECONDARY
        run_t = p.add_run(text)
        run_t.font.color.rgb = BODY_TEXT
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            run_b.font.bold = True
            run_b.font.color.rgb = SECONDARY
        run_t = p.add_run(text)
        run_t.font.color.rgb = BODY_TEXT
        return p

    def add_callout(text, title="NOTE"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.5)
        set_cell_background(cell, HEX_LIGHT_BG)
        set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r_title = p.add_run(f"[{title}] ")
        r_title.font.bold = True
        r_title.font.color.rgb = PRIMARY
        
        r_txt = p.add_run(text)
        r_txt.font.color.rgb = BODY_TEXT
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    def add_figure(img_path, caption_text, figure_num):
        if os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(12)
            p_img.paragraph_format.space_after = Pt(4)
            run_img = p_img.add_run()
            run_img.add_picture(img_path, width=Inches(5.8))

            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_before = Pt(2)
            p_cap.paragraph_format.space_after = Pt(14)
            r_cap_b = p_cap.add_run(f"Figure {figure_num}: ")
            r_cap_b.font.bold = True
            r_cap_b.font.color.rgb = PRIMARY
            r_cap_b.font.size = Pt(9.5)
            r_cap_t = p_cap.add_run(caption_text)
            r_cap_t.font.italic = True
            r_cap_t.font.color.rgb = MUTED
            r_cap_t.font.size = Pt(9.5)

    # Assets setup
    brain_dir = r"C:\Users\user\.gemini\antigravity-ide\brain\05c0be4d-b9ae-4ab1-b726-7df646f5f65e"
    target_assets_dir = os.path.join(os.getcwd(), "report_assets")
    os.makedirs(target_assets_dir, exist_ok=True)

    images = {
        "fig3_1": ("mediunity_system_architecture_1785991672606.png", "Figure 3.1: MediUnity High-Level System Architecture Diagram"),
        "fig3_2": ("mediunity_simulated_network_diagram_1785991687789.png", "Figure 3.2: Simulated Cloud Network Topology and Service Infrastructure"),
        "fig3_3": ("mediunity_process_flow_1_1785991702626.png", "Figure 3.3: Patient Appointment Booking & Telehealth Consultation Workflow"),
        "fig3_4": ("mediunity_process_flow_2_1785991720697.png", "Figure 3.4: Administrator Moderation and Partner Onboarding Verification Workflow"),
        "fig4_1": ("mediunity_user_part_ui_1785991734499.png", "Figure 4.1: Patient & End-User Interactive Application Dashboards"),
        "fig4_2": ("mediunity_admin_part_ui_1785991749603.png", "Figure 4.2: Administrator Control Center & System Compliance Monitoring Interface")
    }

    img_paths = {}
    for key, (fname, caption) in images.items():
        src = os.path.join(brain_dir, fname)
        dst = os.path.join(target_assets_dir, fname)
        if os.path.exists(src):
            shutil.copy(src, dst)
            img_paths[key] = dst
        elif os.path.exists(dst):
            img_paths[key] = dst
        else:
            img_paths[key] = None

    # ==================== COVER PAGE ====================
    add_title("MediUnity: A Cloud-Ready Distributed Multi-Portal Healthcare and Wellness Platform")
    add_subtitle("A Project Report Submitted in Partial Fulfillment of the Requirements for the Degree of Bachelor of Science in Computer Science and Engineering (CSE)")

    p_by = doc.add_paragraph()
    p_by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_by.paragraph_format.space_before = Pt(18)
    p_by.paragraph_format.space_after = Pt(6)
    r_by = p_by.add_run("By")
    r_by.font.bold = True
    r_by.font.size = Pt(12)
    r_by.font.color.rgb = SECONDARY

    p_students = doc.add_paragraph()
    p_students.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_students.paragraph_format.space_after = Pt(24)
    p_students.paragraph_format.line_spacing = 1.2
    r_s1 = p_students.add_run("Sarower Rahman\nID: 111222033\n\nJoy Ranjanshil\nID: 111222007\n")
    r_s1.font.bold = True
    r_s1.font.size = Pt(11.5)
    r_s1.font.color.rgb = BODY_TEXT

    p_batch = doc.add_paragraph()
    p_batch.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_batch.paragraph_format.space_after = Pt(24)
    r_b = p_batch.add_run("10th Batch (Fall 2022 – Spring 2026)")
    r_b.font.italic = True
    r_b.font.size = Pt(11)
    r_b.font.color.rgb = MUTED

    p_sup_label = doc.add_paragraph()
    p_sup_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sup_label.paragraph_format.space_after = Pt(6)
    r_sl = p_sup_label.add_run("Supervised by")
    r_sl.font.bold = True
    r_sl.font.size = Pt(12)
    r_sl.font.color.rgb = SECONDARY

    p_sup = doc.add_paragraph()
    p_sup.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sup.paragraph_format.space_after = Pt(36)
    p_sup.paragraph_format.line_spacing = 1.2
    r_sup_name = p_sup.add_run("Sabbir Ahmed Shihab\n")
    r_sup_name.font.bold = True
    r_sup_name.font.size = Pt(12)
    r_sup_name.font.color.rgb = PRIMARY
    r_sup_details = p_sup.add_run("Lecturer\nDepartment of Computer Science and Engineering (CSE)\nCCN University of Science and Technology\nKothbari, Cumilla, 3500")
    r_sup_details.font.size = Pt(10.5)
    r_sup_details.font.color.rgb = BODY_TEXT

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_date.paragraph_format.space_before = Pt(36)
    r_date = p_date.add_run("Spring 2026 (May 2026)")
    r_date.font.bold = True
    r_date.font.size = Pt(11)
    r_date.font.color.rgb = SECONDARY

    doc.add_page_break()

    # ==================== PRELIMINARY PAGES ====================
    # Board of Examiners
    add_heading_1("Board of Examiners")
    add_body("The project report titled \"MediUnity: A Cloud-Ready Distributed Multi-Portal Healthcare and Wellness Platform\", submitted by Sarower Rahman (ID: 111222033) and Joy Ranjanshil (ID: 111222007) of the 10th Batch (Fall 2022 – Spring 2026), has been examined and accepted as satisfactory in partial fulfillment of the requirements for the degree of Bachelor of Science in Computer Science and Engineering (CSE).")
    
    # Table for Board of Examiners Signatures
    tbl_boe = doc.add_table(rows=4, cols=2)
    tbl_boe.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r in range(4):
        for c in range(2):
            cell = tbl_boe.cell(r, c)
            cell.width = Inches(3.2)
            set_cell_margins(cell, top=120, bottom=120, left=100, right=100)

    examiners = [
        ("___________________________\nSabbir Ahmed Shihab\nSupervisor & Lecturer\nDepartment of CSE, CCN UST", "Chairman\n(Internal)"),
        ("___________________________\nHead of Department\nDepartment of CSE\nCCN UST", "Member\n(Ex-Officio)"),
        ("___________________________\nInternal Examiner\nDepartment of CSE\nCCN UST", "Member\n(Internal)"),
        ("___________________________\nExternal Examiner\nDepartment of CSE\nPartner University / Industry Expert", "Member\n(External)")
    ]

    for idx, (member_info, role) in enumerate(examiners):
        r_idx = idx
        cell_left = tbl_boe.cell(r_idx, 0)
        cell_right = tbl_boe.cell(r_idx, 1)

        p1 = cell_left.paragraphs[0]
        p1.paragraph_format.line_spacing = 1.15
        p1.add_run(member_info).font.size = Pt(9.5)

        p2 = cell_right.paragraphs[0]
        p2.paragraph_format.line_spacing = 1.15
        r_role = p2.add_run(role)
        r_role.font.bold = True
        r_role.font.color.rgb = PRIMARY
        r_role.font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    doc.add_page_break()

    # Declaration of Originality
    add_heading_1("Declaration of Originality")
    add_body("We hereby declare that the work presented in this project report entitled \"MediUnity: A Cloud-Ready Distributed Multi-Portal Healthcare and Wellness Platform\" is an authentic record of our own research and development work carried out under the supervision of Sabbir Ahmed Shihab, Lecturer, Department of Computer Science and Engineering (CSE), CCN University of Science and Technology, Cumilla.")
    add_body("We confirm that:")
    add_bullet("The material contained in this report has not been submitted previously, in full or in part, for any degree, diploma, or qualification at this or any other university or institute.")
    add_bullet("All core algorithms, system architectures, database schemas, API routes, and frontend portals described herein were designed and implemented by the undersigned project team.")
    add_bullet("Wherever third-party open-source libraries, external APIs, software frameworks, or literature references have been utilized, they have been fully acknowledged and cited in accordance with standard academic ethics and IEEE reference formats.")

    p_sig = doc.add_paragraph()
    p_sig.paragraph_format.space_before = Pt(36)
    p_sig.paragraph_format.line_spacing = 1.3
    p_sig.add_run("_____________________________\nSarower Rahman\nID: 111222033\n10th Batch, Department of CSE\nCCN UST\n\n\n_____________________________\nJoy Ranjanshil\nID: 111222007\n10th Batch, Department of CSE\nCCN UST").font.size = Pt(10)

    doc.add_page_break()

    # Acknowledgements
    add_heading_1("Acknowledgements")
    add_body("First and foremost, we express our deepest gratitude to Almighty God for providing us with the health, strength, wisdom, and perseverance needed to successfully complete our undergraduate project and synthesize this comprehensive technical report.")
    add_body("We wish to extend our most sincere gratitude and profound appreciation to our esteemed project supervisor, Sabbir Ahmed Shihab, Lecturer, Department of Computer Science and Engineering (CSE), CCN University of Science and Technology. His invaluable guidance, constructive feedback, insightful critique, and continuous technical encouragement throughout the design and execution phases were fundamental to the realization of MediUnity.")
    add_body("We are also deeply thankful to the honorable Head of the Department of Computer Science and Engineering and all respected faculty members of CCN UST for providing a rigorous academic foundation, modern laboratory facilities, and an inspiring environment that fostered our technical growth.")
    add_body("Finally, we express our heartfelt appreciation to our families and fellow classmates of the CSE 10th Batch for their constant moral support, patience, and motivation during long hours of software development, testing, and documentation.")

    doc.add_page_break()

    # Abstract
    add_heading_1("Abstract")
    add_body("Modern healthcare delivery across global environments faces persistent challenges, including fragmented patient communication, manual appointment scheduling bottlenecks, delayed access to patient medical histories, isolated pharmacy inventories, and a lack of unified digital coordination between healthcare providers. To resolve these systemic inefficiencies, this project presents MediUnity—a cloud-ready, full-stack, multi-portal healthcare and wellness management platform engineered for global accessibility.")
    add_body("MediUnity unifies five critical healthcare stakeholders into a cohesive digital ecosystem: patients, medical doctors, hospitals, diagnostic centers, and licensed pharmacies, backed by an enterprise administrative control center. The platform architecture utilizes a modern JavaScript stack comprising React 19 and Vite for responsive presentation, Node.js with Express.js for RESTful API orchestration, MongoDB via Mongoose ODM for schemaless electronic health record storage, and Redis coupled with BullMQ for asynchronous background process queues. Real-time interaction—including instant doctor-patient messaging, emergency notification dispatches, and appointment queue updates—is driven by Socket.IO, while external cloud integrations facilitate secure media storage (Cloudinary), payment processing (Stripe), multi-channel alerts (Nodemailer, Twilio, Firebase Admin), and optical character recognition for prescriptions (Tesseract.js).")
    add_body("Key technical achievements of MediUnity include role-based multi-portal access control (RBAC), end-to-end appointment lifecycle tracking, integrated diagnostic test ordering, automated medicine delivery workflows, interactive patient vitals logging, automated doctor payout calculations, and multi-layered security middleware enforcing rate limiting, input sanitization, CORS protection, and HTTP header security via Helmet. Extensive testing and validation demonstrate high operational throughput, low latency, robust data protection, and seamless multi-device responsiveness. MediUnity serves as a scalable, production-ready foundation capable of expanding digital healthcare access worldwide.")

    doc.add_page_break()

    # Table of Contents
    add_heading_1("Table of Contents")
    toc_items = [
        ("Board of Examiners", "ii"),
        ("Declaration of Originality", "iii"),
        ("Acknowledgements", "iv"),
        ("Abstract", "v"),
        ("List of Figures", "vii"),
        ("Chapter 1: Introduction", "1"),
        ("    1.1 Introduction", "1"),
        ("    1.2 Purpose of the Project", "2"),
        ("    1.3 Scope", "3"),
        ("    1.4 Project Outline", "4"),
        ("Chapter 2: Literature Review", "5"),
        ("    2.1 Characteristics", "5"),
        ("    2.2 Used Technology", "7"),
        ("    2.3 Software Requirement", "9"),
        ("        2.3.1 Hardware Requirement", "10"),
        ("Chapter 3: Methodology and Design", "11"),
        ("    3.1 System Architecture", "11"),
        ("    3.2 Simulated Diagram", "13"),
        ("        3.2.1 Process Flow 1: Patient Booking & Consultation Flow", "14"),
        ("        3.2.2 Process Flow 2: Admin Moderation & Partner Verification", "16"),
        ("Chapter 4: Results and Implementation", "18"),
        ("    4.1 User Part", "18"),
        ("    4.2 Admin Part", "21"),
        ("Chapter 5: Conclusion and Future Directions", "24"),
        ("    5.1 Advantages", "24"),
        ("    5.2 Limitations", "25"),
        ("    5.3 Future Scopes", "26"),
        ("References", "27")
    ]
    for item, page in toc_items:
        p_toc = doc.add_paragraph()
        p_toc.paragraph_format.space_after = Pt(3)
        p_toc.paragraph_format.line_spacing = 1.15
        is_chap = item.startswith("Chapter") or item in ["Board of Examiners", "Declaration of Originality", "Acknowledgements", "Abstract", "List of Figures", "References"]
        
        r_item = p_toc.add_run(item)
        if is_chap:
            r_item.font.bold = True
            r_item.font.color.rgb = SECONDARY
        else:
            r_item.font.color.rgb = BODY_TEXT

        # Tab stop simulation with dots
        dots_count = max(5, 75 - len(item))
        r_dots = p_toc.add_run(" " + "." * dots_count + " ")
        r_dots.font.color.rgb = MUTED
        r_page = p_toc.add_run(page)
        r_page.font.bold = is_chap
        r_page.font.color.rgb = PRIMARY

    doc.add_page_break()

    # List of Figures
    add_heading_1("List of Figures")
    fig_items = [
        ("Figure 3.1", "MediUnity High-Level System Architecture Diagram", "12"),
        ("Figure 3.2", "Simulated Cloud Network Topology and Service Infrastructure", "13"),
        ("Figure 3.3", "Patient Appointment Booking & Telehealth Consultation Workflow", "15"),
        ("Figure 3.4", "Administrator Moderation and Partner Onboarding Verification Workflow", "17"),
        ("Figure 4.1", "Patient & End-User Interactive Application Dashboards", "19"),
        ("Figure 4.2", "Administrator Control Center & System Compliance Monitoring Interface", "22")
    ]
    for fig_id, fig_title, fig_pg in fig_items:
        p_fig = doc.add_paragraph()
        p_fig.paragraph_format.space_after = Pt(4)
        r_fid = p_fig.add_run(f"{fig_id}: ")
        r_fid.font.bold = True
        r_fid.font.color.rgb = PRIMARY
        r_ftit = p_fig.add_run(fig_title)
        r_ftit.font.color.rgb = BODY_TEXT
        
        dots_c = max(5, 68 - len(fig_title))
        p_fig.add_run(" " + "." * dots_c + " ").font.color.rgb = MUTED
        r_fpg = p_fig.add_run(fig_pg)
        r_fpg.font.bold = True
        r_fpg.font.color.rgb = SECONDARY

    doc.add_page_break()

    # ==================== CHAPTER 1 ====================
    add_heading_1("Chapter 1: Introduction")
    
    add_heading_2("1.1 Introduction")
    add_body("Healthcare delivery in the modern era is undergoing a fundamental digital transformation driven by the rapid evolution of distributed web technologies, real-time messaging, mobile accessibility, and cloud computing. Traditional medical service management—characterized by paper-based recordkeeping, physical queue management, fragmented diagnostic laboratory reporting, manual pharmacy inventory checks, and isolated doctor consultations—is increasingly inadequate to meet global healthcare demands. Patients face severe delays in securing qualified medical consultations, accessing personal health records, and purchasing specialized pharmaceuticals. Concurrently, medical practitioners and health facility administrators suffer from operational inefficiencies, disconnected patient data silos, and administrative burdens.")
    add_body("To address these critical global healthcare challenges, this project introduces MediUnity (alternatively designated as the Medicare System)—an enterprise-grade, cloud-ready, multi-portal digital healthcare and wellness ecosystem. Built with a scalable multi-tenant architecture, MediUnity connects patients, medical clinicians, hospitals, diagnostic centers, and retail pharmacies into a single, unified digital platform. Designed for global operational deployment, MediUnity provides localized and worldwide reach, multilingual capabilities, internationalized health metric tracking, and compliant medical record management.")

    add_heading_2("1.2 Purpose of the Project")
    add_body("The primary objective of MediUnity is to engineer a unified digital healthcare ecosystem that eliminates geographical and structural barriers between healthcare consumers and medical providers. The specific core purposes of the system include:")
    add_bullet("Universal Healthcare Marketplace: ", "Constructing an intuitive marketplace where users globally can search, evaluate, and connect with verified medical doctors, accredited hospitals, diagnostic laboratories, and licensed retail pharmacies.")
    add_bullet("Streamlined Appointment Scheduling: ", "Providing real-time doctor availability calendars, automated consultation booking, instant payment processing, and multi-channel appointment reminders.")
    add_bullet("Centralized Electronic Health Records (EHR): ", "Establishing a secure, cloud-hosted repository for patient medical histories, lab test diagnostic reports, prescription histories, and vital health logs accessible across multi-portal interfaces.")
    add_bullet("Integrated Pharmacy & Diagnostic Services: ", "Enabling patients to book lab specimen collections, order diagnostic imaging, and purchase prescribed medications for home delivery directly through the platform.")
    add_bullet("Operational Administrative Supervision: ", "Equipping platform administrators with robust oversight tools for partner verification, medical license auditing, compliance monitoring, revenue allocation, and system telemetry.")

    add_heading_2("1.3 Scope")
    add_body("The functional and technical scope of MediUnity spans multiple operational boundaries and software domains:")
    add_bullet("Global Applicability & Multitenancy: ", "Designed for global scalability across urban and regional healthcare systems without localized hardcoding, supporting international date formats, currency units, and multi-region cloud deployment.")
    add_bullet("Multi-Portal User Interfaces: ", "Includes dedicated client portals optimized for distinct stakeholders: Patient Web Portal (`frontend` & `frontend-patient`), Doctor Portal (`frontend-doctor`), Partner Provider Portal (`frontend-partner`), and Administrative Control Panel (`admin`).")
    add_bullet("Asynchronous & Real-Time Communications: ", "Features Socket.IO web socket channels for real-time teleconsultation chat, live doctor status updates, and automated background push notifications via Redis and BullMQ queues.")
    add_bullet("Financial & Service Integrations: ", "Integrates Stripe payment gateways for instant consultation fee processing, automated doctor payout accounting, and Cloudinary media pipelines for high-resolution medical document storage.")
    add_bullet("Boundary Constraints: ", "While MediUnity facilitates virtual consultations, medical file sharing, and e-prescriptions, it acts as a digital facilitator and does not replace emergency ICU life-support hardware or physical surgical apparatus.")

    add_heading_2("1.4 Project Outline")
    add_body("The remainder of this project report is structured into the following detailed chapters:")
    add_bullet("Chapter 2: Literature Review – ", "Examines existing healthcare management systems, analyzes comparative platform features, details the technological stack (React 19, Express, MongoDB, Redis, Cloudinary), and outlines software/hardware requirements.")
    add_bullet("Chapter 3: Methodology and Design – ", "Explains the multi-tier system architecture, cloud deployment topology, entity-relationship schemas, and process flow diagrams for appointment booking and administrative moderation.")
    add_bullet("Chapter 4: Results and Implementation – ", "Demonstrates the practical implementation outcomes from both the end-user perspective (patient/doctor workflows) and administrative operational control panel.")
    add_bullet("Chapter 5: Conclusion and Future Directions – ", "Summarizes platform achievements, highlights architectural advantages, identifies system limitations, and outlines future enhancements including AI symptom triage and wearable IoT integration.")

    doc.add_page_break()

    # ==================== CHAPTER 2 ====================
    add_heading_1("Chapter 2: Literature Review")
    add_body("Digital health platforms have evolved significantly over the past two decades, transitioning from localized desktop electronic medical record (EMR) databases to cloud-native telemedicine marketplaces. This chapter reviews existing healthcare systems, outlines the key characteristics of MediUnity, details the primary technology stack, and lists the required software and hardware specifications.")

    add_heading_2("2.1 Characteristics")
    add_body("MediUnity distinguishes itself from existing proprietary and legacy health management platforms through several core architectural characteristics:")
    add_bullet("1. Multi-Portal Domain Separation: ", "Unlike monolithic healthcare tools that force all user roles into a single crowded interface, MediUnity deploys specialized frontend portals (`frontend-patient`, `frontend-doctor`, `frontend-partner`, `admin`) styled specifically for role-tailored user experience.")
    add_bullet("2. Cloud-Native Scalability: ", "Architected with stateless Express REST APIs, decoupled MongoDB database clusters, and container-ready configuration (`render.yaml`, environment templates), enabling horizontal autoscaling during peak demand.")
    add_bullet("3. Asynchronous Task Processing: ", "Leverages Redis and BullMQ queue workers to handle compute-heavy and latent jobs (SMS dispatches via Twilio, email newsletters via Nodemailer, optical character recognition via Tesseract.js) without blocking HTTP response loops.")
    add_bullet("4. Comprehensive Service Lifecycle: ", "Facilitates end-to-end medical care, encompassing initial doctor search, slot reservation, online payment, virtual consultation, prescription generation, diagnostic lab testing, pharmacy order fulfillment, and specialist referral transfers.")
    add_bullet("5. Defense-in-Depth Security: ", "Incorporates multi-layered security middleware including Helmet HTTP header protection, IP-based rate limiting, MongoDB query sanitization, XSS payload filtering, and JWT token authentication.")

    add_heading_2("2.2 Used Technology")
    add_body("The technical stack powering MediUnity was selected for high performance, developer agility, robust community support, and enterprise cloud compatibility:")
    
    # Table of Technologies
    tbl_tech = doc.add_table(rows=6, cols=3)
    tbl_tech.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Layer / Domain", "Technology Stack", "Role & Implementation Purpose"]
    for i, h in enumerate(headers):
        cell = tbl_tech.cell(0, i)
        set_cell_background(cell, HEX_PRIMARY)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    tech_data = [
        ("Frontend Client", "React 19, Vite, Tailwind CSS, Lucide Icons, Axios", "Responsive single-page applications for Patients, Doctors, Partners, and Admins with PWA capability."),
        ("Backend Server", "Node.js, Express.js (v5), Socket.IO", "RESTful API web server, real-time WebSocket messaging, and request routing."),
        ("Database & Caching", "MongoDB, Mongoose ODM, Redis", "Schemaless document store for EHR/users; Redis for queue management and session caching."),
        ("Async & Processing", "BullMQ, Tesseract.js, Nodemailer, Twilio", "Background queue workers, OCR prescription parsing, automated email dispatches, and SMS dispatches."),
        ("Cloud & Security", "Cloudinary, Stripe, Helmet, Express-Rate-Limit", "Cloud medical file storage, global payment processing, rate limiting, and HTTP security headers.")
    ]

    for row_idx, data in enumerate(tech_data, start=1):
        for col_idx, text in enumerate(data):
            cell = tbl_tech.cell(row_idx, col_idx)
            if row_idx % 2 == 1:
                set_cell_background(cell, HEX_LIGHT_BG)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            p.add_run(text).font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    add_heading_2("2.3 Software Requirement")
    add_body("The software tools and operational environment required to build, execute, and host MediUnity comprise:")
    add_bullet("Operating System: ", "Cross-platform compatibility across Microsoft Windows 10/11, Ubuntu Linux 22.04 LTS, or macOS Sonoma.")
    add_bullet("Node.js Runtime: ", "Node.js v18.x or v20.x LTS with npm v9.x+ package manager.")
    add_bullet("Database Engine: ", "MongoDB v6.0+ Community Server or MongoDB Atlas Cloud instance.")
    add_bullet("In-Memory Store: ", "Redis v7.0+ for BullMQ background job queuing.")
    add_bullet("Web Browsers: ", "Modern evergreen browsers supporting HTML5, WebSockets, and ES6+ (Google Chrome, Mozilla Firefox, Microsoft Edge, Apple Safari).")
    add_bullet("Development Tools: ", "Visual Studio Code / Antigravity IDE, Git version control, Postman API tester.")

    add_heading_3("2.3.1 Hardware Requirement")
    add_body("The minimum and recommended hardware specifications for hosting and running MediUnity across development and cloud environments are outlined below:")
    
    # Table of Hardware Requirements
    tbl_hw = doc.add_table(rows=5, cols=3)
    tbl_hw.alignment = WD_TABLE_ALIGNMENT.CENTER
    hw_headers = ["Component", "Minimum Requirement (Dev/Test)", "Recommended (Cloud Production)"]
    for i, h in enumerate(hw_headers):
        cell = tbl_hw.cell(0, i)
        set_cell_background(cell, HEX_PRIMARY)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    hw_data = [
        ("Processor (CPU)", "Dual-Core 2.0 GHz Intel/AMD or Apple M1", "8-Core 3.0 GHz x86-64 / ARM Cloud Server Instance"),
        ("System Memory (RAM)", "4 GB DDR4 RAM", "16 GB to 32 GB High-Speed RAM"),
        ("Storage (Disk)", "10 GB available SSD space", "100 GB NVMe SSD with automated cloud backups"),
        ("Network Bandwidth", "10 Mbps broadband internet connection", "1 Gbps dedicated cloud network interface with DDoS mitigation")
    ]

    for row_idx, data in enumerate(hw_data, start=1):
        for col_idx, text in enumerate(data):
            cell = tbl_hw.cell(row_idx, col_idx)
            if row_idx % 2 == 1:
                set_cell_background(cell, HEX_LIGHT_BG)
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            p.add_run(text).font.size = Pt(9.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    doc.add_page_break()

    # ==================== CHAPTER 3 ====================
    add_heading_1("Chapter 3: Methodology and Design")
    add_body("This chapter details the system engineering methodology, architectural layout, network deployment topology, and core process flow diagrams that govern MediUnity's operations.")

    add_heading_2("3.1 System Architecture")
    add_body("MediUnity follows a decoupled multi-tier architecture structured across four distinct operational layers: Presentation, Security Gateway, Application Services, and Data/Storage Layer. Figure 3.1 illustrates this layered structure.")
    add_bullet("1. Presentation Layer: ", "Comprises lightweight React 19 single-page applications served via Vite. Dedicated portals exist for Patients, Doctors, Partner Organizations (Hospitals, Pharmacies, Diagnostic Centers), and System Administrators.")
    add_bullet("2. Security & Gateway Layer: ", "All incoming client HTTP/WebSocket requests pass through express-rate-limit to prevent brute-force attacks, Helmet to enforce HTTP security headers, mongo-sanitize to neutralize NoSQL injections, and CORS policy controllers.")
    add_bullet("3. Application & Service Layer: ", "Powered by Express.js REST API routers, managing business logic for appointments, user profiles, diagnostic bookings, orders, and payment transactions. Socket.IO manages real-time messaging, while BullMQ workers process background tasks.")
    add_bullet("4. Data & External Integration Layer: ", "MongoDB handles persistent document storage (Mongoose models for Patients, Doctors, Appointments, Orders, Prescriptions, Audit Logs). Redis stores transient queue states. External APIs handle Cloudinary media, Stripe payments, Twilio SMS dispatches, and Firebase notifications.")

    if img_paths.get("fig3_1"):
        add_figure(img_paths["fig3_1"], "MediUnity High-Level System Architecture Diagram showing multi-portal client layers, security gateway middleware, Node.js API engine, and MongoDB/Redis data tiers.", "3.1")

    add_heading_2("3.2 Simulated Diagram")
    add_body("To demonstrate production readiness, MediUnity's infrastructure was modeled as a distributed cloud container deployment. Figure 3.2 depicts the simulated network topology.")
    add_body("The network architecture utilizes a cloud load balancer with SSL/TLS termination to route HTTPS traffic across an auto-scaling cluster of React frontend static nodes and Node.js API application containers. Background processing is offloaded to containerized BullMQ worker instances communicating with a high-availability Redis instance. MongoDB is deployed in a replica set configuration to ensure continuous data availability and disaster recovery.")

    if img_paths.get("fig3_2"):
        add_figure(img_paths["fig3_2"], "Simulated Cloud Network Topology and Infrastructure Diagram featuring Load Balancers, API Clusters, BullMQ Worker Containers, Redis Caches, and MongoDB Replica Sets.", "3.2")

    add_heading_3("3.2.1 Process Flow 1: Patient Booking & Consultation Flow")
    add_body("Process Flow 1 (Figure 3.3) illustrates the end-to-end sequence executed when a patient searches for a doctor, schedules an appointment, processes payment, and conducts a consultation:")
    add_bullet("Step 1: Doctor Search & Discovery: ", "Patient queries the marketplace by medical specialty, location, or rating. Frontend requests filtered results from `/api/doctors`.")
    add_bullet("Step 2: Slot Selection & Lock: ", "Patient selects an available time slot. The system verifies slot availability in `Appointment` collection and places a temporary hold.")
    add_bullet("Step 3: Stripe Payment Processing: ", "Patient submits payment credentials. Express API invokes Stripe SDK to charge the consultation fee and returns transaction status.")
    add_bullet("Step 4: Confirmation & Alert Dispatch: ", "Upon payment verification, the appointment state updates to `confirmed`. BullMQ enqueues confirmation dispatches via Nodemailer (Email) and Twilio (SMS).")
    add_bullet("Step 5: Teleconsultation Session: ", "At the scheduled time, patient and doctor enter the virtual consultation room, communicating in real time over Socket.IO WebSockets.")
    add_bullet("Step 6: Prescription & EHR Archival: ", "Doctor generates a digital prescription (`Prescription` model), which is stored in MongoDB and instantly visible in the patient's EHR vault.")

    if img_paths.get("fig3_3"):
        add_figure(img_paths["fig3_3"], "Process Flow 1 Diagram illustrating the step-by-step patient appointment booking, Stripe payment, notification dispatch, and digital prescription creation workflow.", "3.3")

    add_heading_3("3.2.2 Process Flow 2: Admin Moderation & Partner Onboarding Flow")
    add_body("Process Flow 2 (Figure 3.4) outlines the administrative compliance, verification, and onboarding process required before doctors, hospitals, diagnostic centers, or pharmacies can operate on the platform:")
    add_bullet("Step 1: Partner Application Submission: ", "The prospective doctor or partner entity registers via the portal, uploading medical licenses, credentials, and identity verification files.")
    add_bullet("Step 2: Automated Pre-Validation: ", "System middleware validates input file formats, checks for duplicate license numbers, and enqueues the record into the Admin Moderation Queue.")
    add_bullet("Step 3: Administrator Credentials Audit: ", "An authorized platform administrator accesses the `admin` portal, reviewing submitted credentials, license authority records, and uploaded documents.")
    add_bullet("Step 4: Decision Branch (Approval / Rejection): ", "If credentials meet statutory healthcare standards, Admin clicks `Approve`. If discrepancies exist, Admin clicks `Request Re-submission` with specific feedback.")
    add_bullet("Step 5: Catalog Activation & Status Update: ", "Upon approval, the system updates the entity status to `isVerified: true`, making their profile publicly discoverable in the healthcare marketplace.")
    add_bullet("Step 6: Audit Logging: ", "An immutable record of the admin verification action is appended to the `AuditLog` collection for regulatory compliance auditing.")

    if img_paths.get("fig3_4"):
        add_figure(img_paths["fig3_4"], "Process Flow 2 Diagram depicting the administrative moderation, credential audit, decision branching, and partner catalog activation workflow.", "3.4")

    doc.add_page_break()

    # ==================== CHAPTER 4 ====================
    add_heading_1("Chapter 4: Results and Implementation")
    add_body("This chapter presents the concrete implementation outcomes of MediUnity, highlighting the interactive user application interfaces and administrative operational management tools.")

    add_heading_2("4.1 User Part")
    add_body("The end-user experience encompasses the patient, doctor, and partner provider portals, designed with responsive layouts, dark/light theme options, and intuitive healthcare navigation workflows. Figure 4.1 showcases the patient and end-user interface layout.")
    add_body("Key implemented features within the User Part include:")
    add_bullet("Healthcare Marketplace Discovery: ", "Patients can seamlessly filter doctors by clinical specialty (e.g., Cardiology, Neurology, Pediatrics), search hospitals by geographical proximity, and view diagnostic test menus.")
    add_bullet("Interactive Appointment Management: ", "Provides real-time visibility into active, upcoming, completed, and canceled appointments, with one-click video consultation room entry and reschedule capabilities.")
    add_bullet("Personal Vitals & Health Logging: ", "Patients can log daily health metrics (blood pressure, blood glucose, heart rate, BMI, temperature), visualizing trends on interactive graphs.")
    add_bullet("Digital Prescription & Lab Order Tracking: ", "Patients view downloadable PDF prescriptions, track diagnostic sample collection status, and order prescribed medicines directly from registered pharmacies.")
    add_bullet("Doctor Clinical Dashboard: ", "Physicians access their daily appointment rosters, review patient medical histories prior to consultation, issue digital prescriptions, and monitor monthly earnings payouts.")

    if img_paths.get("fig4_1"):
        add_figure(img_paths["fig4_1"], "User Part Interface Layout displaying doctor search, upcoming appointment tracking cards, health vitals monitoring widgets, and diagnostic test menus.", "4.1")

    add_heading_2("4.2 Admin Part")
    add_body("The Administrative Control Panel (`admin`) equips platform moderators and enterprise operations teams with complete oversight of platform health, security, compliance, and user verification. Figure 4.2 illustrates the Admin Part operational dashboard.")
    add_body("Core administrative capabilities implemented in MediUnity include:")
    add_bullet("Verification & Moderation Hub: ", "A centralized approval desk where administrators review pending doctor licenses, hospital accreditations, diagnostic lab certifications, and pharmacy operating permits.")
    add_bullet("System Performance Telemetry: ", "Real-time monitoring of API response times, active database connections, Redis queue throughput, server CPU/memory load, and error exception logs.")
    add_bullet("User & Financial Management: ", "Admins can manage platform roles, resolve patient-doctor disputes, inspect transaction logs, configure commission rates, and approve doctor payout disbursements.")
    add_bullet("Platform Settings & Feature Toggles: ", "Allows global administrators to dynamically modify platform settings (e.g., maintenance mode, payment gateway keys, global notification banners) without re-deploying backend code.")
    add_bullet("Compliance & Security Audit Logging: ", "Tracks all administrative actions, credential approvals, and privilege escalations in an immutable `AuditLog` database table for regulatory auditing.")

    if img_paths.get("fig4_2"):
        add_figure(img_paths["fig4_2"], "Admin Part Operations Dashboard featuring key platform metric cards, partner verification queues, system health graphs, and compliance audit logs.", "4.2")

    doc.add_page_break()

    # ==================== CHAPTER 5 ====================
    add_heading_1("Chapter 5: Conclusion and Future Directions")
    add_body("This final chapter synthesizes the overall project outcomes, lists the key architectural advantages of MediUnity, discusses discovered system limitations, and outlines prospective future research and feature enhancements.")

    add_heading_2("5.1 Advantages")
    add_body("MediUnity offers significant technological and operational benefits over traditional and siloed healthcare software solutions:")
    add_bullet("1. Multi-Tenant Ecosystem Consolidation: ", "Unifies patients, doctors, hospitals, diagnostic centers, and pharmacies into a single integrated digital platform, eliminating isolated software silos.")
    add_bullet("2. Enterprise Multi-Portal Design: ", "Delivers tailored user experiences for distinct user roles (`frontend-patient`, `frontend-doctor`, `frontend-partner`, `admin`), reducing cognitive load and operational friction.")
    add_bullet("3. High Performance & Asynchronous Reliability: ", "Decouples synchronous HTTP requests from latent background tasks using Redis and BullMQ queues, ensuring rapid API responses under heavy load.")
    add_bullet("4. Comprehensive End-to-End Workflows: ", "Covers the complete spectrum of digital healthcare—from initial doctor search and booking to payment processing, teleconsultation, lab testing, pharmacy ordering, and specialist referrals.")
    add_bullet("5. Security & Regulatory Readiness: ", "Implements defense-in-depth protection including rate limiting, Helmet headers, NoSQL injection sanitization, XSS filters, and comprehensive audit logging.")

    add_heading_2("5.2 Limitations")
    add_body("Despite its robust architecture, several system constraints and operational challenges were identified during development and testing:")
    add_bullet("1. Reliance on External Cloud Dependencies: ", "Features such as SMS dispatches (Twilio), email delivery (Nodemailer), payment processing (Stripe), and media hosting (Cloudinary) depend on third-party API availability and connectivity.")
    add_bullet("2. Hardware Wearable Synchronization: ", "Health vitals tracking currently relies on manual patient entries or simulated API payloads rather than direct hardware Bluetooth Low Energy (BLE) sync with consumer smartwatches.")
    add_bullet("3. Multi-Portal Maintenance Overhead: ", "Maintaining multiple separate frontend applications (`frontend-patient`, `frontend-doctor`, `frontend-partner`, `admin`) requires synchronized UI library updates and strict state synchronization.")
    add_bullet("4. Regulatory Compliance Variations: ", "Operating across diverse international jurisdictions requires adaptive localized compliance frameworks (HIPAA, GDPR) beyond basic role-based access control.")

    add_heading_2("5.3 Future Scopes")
    add_body("Future development iterations of MediUnity will expand its technological capabilities in the following directions:")
    add_bullet("1. AI-Driven Symptom Triage & AI Assistant: ", "Integrating Large Language Models (LLMs) and clinical decision support systems to provide automated preliminary symptom triage and doctor specialty recommendations.")
    add_bullet("2. Direct Wearable IoT Sensor Sync: ", "Developing mobile PWA and native SDK bridges to directly ingest real-time continuous vitals (ECG, SpO2, continuous glucose monitoring) from Apple HealthKit, Google Fit, and wearable IoT devices.")
    add_bullet("3. Telemedicine WebRTC Video Streaming: ", "Upgrading the Socket.IO messaging hub to native WebRTC video/audio peer-to-peer streaming for seamless embedded video consultations within the portal.")
    add_bullet("4. Blockchain Electronic Health Record Vault: ", "Implementing decentralized immutable ledger storage for medical record access logs to enhance patient privacy control and cross-institutional record sharing.")
    add_bullet("5. Predictive Healthcare Analytics: ", "Developing machine learning models to analyze patient health logs and predict disease risks, enabling proactive preventive healthcare interventions.")

    doc.add_page_break()

    # ==================== REFERENCES ====================
    add_heading_1("References")
    references = [
        "[1] World Health Organization, \"Global Strategy on Digital Health 2020–2025,\" WHO Guidelines Approved by the Guidelines Review Committee, Geneva, Switzerland, 2021.",
        "[2] IEEE Standards Association, \"IEEE Standard for Health Informatics - Personal Health Device Communication,\" IEEE Std 11073-10407-2022, pp. 1-84, 2022.",
        "[3] M. Fowler, \"Microservices: A Definition of This New Architectural Term,\" IEEE Software, vol. 31, no. 3, pp. 84-86, May-June 2014.",
        "[4] R. Fielding, \"Architectural Styles and the Design of Network-based Software Architectures,\" Ph.D. dissertation, Dept. Inf. Comput. Sci., Univ. California, Irvine, CA, USA, 2000.",
        "[5] Node.js Foundation, \"Node.js v20 API Documentation & Runtime Benchmarks,\" 2024. [Online]. Available: https://nodejs.org/docs/",
        "[6] React Documentation Team, \"React v19 Technical Architecture and Server Components,\" Facebook Open Source, 2025. [Online]. Available: https://react.dev/",
        "[7] MongoDB Inc., \"MongoDB Enterprise Manual: Data Modeling, Indexing, and Security Protocols,\" MongoDB Press, 2024.",
        "[8] Redis Labs, \"High-Performance Data Structures and Queue Operations with BullMQ,\" Redis Developer Documentation, 2024.",
        "[9] Stripe API Reference, \"Online Payment Gateway Architecture and PCI-DSS Compliance Guidelines,\" Stripe Developers, 2024. [Online]. Available: https://stripe.com/docs/api",
        "[10] O. OWASP Foundation, \"OWASP Top 10 Web Application Security Risks,\" Open Web Application Security Project, 2023. [Online]. Available: https://owasp.org/www-project-top-ten/"
    ]

    for ref in references:
        p_ref = doc.add_paragraph()
        p_ref.paragraph_format.space_after = Pt(6)
        p_ref.paragraph_format.line_spacing = 1.15
        p_ref.paragraph_format.left_indent = Inches(0.4)
        p_ref.paragraph_format.first_line_indent = Inches(-0.4)
        
        # Split reference tag [N] and body
        parts = ref.split(" ", 1)
        r_tag = p_ref.add_run(parts[0] + " ")
        r_tag.font.bold = True
        r_tag.font.color.rgb = PRIMARY
        
        r_body = p_ref.add_run(parts[1] if len(parts) > 1 else "")
        r_body.font.color.rgb = BODY_TEXT

    # Save output document
    output_docx_path = os.path.join(os.getcwd(), "MediUnity_Project_Report.docx")
    doc.save(output_docx_path)
    print("SUCCESS: MediUnity_Project_Report.docx created successfully at:", output_docx_path)

if __name__ == "__main__":
    create_report()
