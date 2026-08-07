import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_callout_box(doc, text_paragraphs, title="FINANCIAL HIGHLIGHT"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, "F1F5F9")
    set_cell_margins(cell, top=160, bottom=160, left=200, right=200)
    
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>\n'
        f'  <w:left w:val="single" w:sz="36" w:space="0" w:color="0F766E"/>\n'
        f'  <w:top w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'  <w:bottom w:val="none"/>\n'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_before = Pt(2)
    p0.paragraph_format.space_after = Pt(4)
    run_t = p0.add_run(f"💰 {title}")
    run_t.font.name = 'Calibri'
    run_t.font.size = Pt(11)
    run_t.font.bold = True
    run_t.font.color.rgb = RGBColor(15, 118, 110)
    
    for text in text_paragraphs:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(51, 65, 85)

def build_monetization_document(output_path):
    doc = docx.Document()
    
    # Page setup - Standard 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)

    # Color Palette
    PRIMARY = RGBColor(15, 118, 110)     # Deep Teal #0F766E
    SECONDARY = RGBColor(15, 23, 42)    # Slate Navy #0F172A
    BODY_TEXT = RGBColor(51, 65, 85)    # Slate 700 #334155
    MUTED = RGBColor(100, 116, 139)     # Slate 500 #64748B
    DARK_BLUE = RGBColor(30, 58, 138)   # Deep Blue #1E3A8A

    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = BODY_TEXT

    def add_doc_header():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        r_meta = p.add_run("MEDIUNITY — HEALTHCARE PLATFORM MONETIZATION REPORT")
        r_meta.font.size = Pt(9.5)
        r_meta.font.bold = True
        r_meta.font.color.rgb = MUTED
        
        p_main = doc.add_paragraph()
        p_main.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_main.paragraph_format.space_before = Pt(4)
        p_main.paragraph_format.space_after = Pt(6)
        r_title = p_main.add_run("BUSINESS MONETIZATION STRATEGY & REVENUE MODEL")
        r_title.font.size = Pt(22)
        r_title.font.bold = True
        r_title.font.color.rgb = PRIMARY
        
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub.paragraph_format.space_before = Pt(0)
        p_sub.paragraph_format.space_after = Pt(20)
        r_sub = p_sub.add_run("Comprehensive Breakdown of Platform Monetization Channels, Commission Mechanics, SaaS Subscriptions & Financial Projections")
        r_sub.font.size = Pt(11.5)
        r_sub.font.italic = True
        r_sub.font.color.rgb = SECONDARY

        p_div = doc.add_paragraph()
        p_div.paragraph_format.space_after = Pt(18)
        p_div_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="1" w:color="0F766E"/></w:pBdr>')
        p_div._element.get_or_add_pPr().append(p_div_border)

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(22)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = PRIMARY
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = SECONDARY
        return p

    def add_p(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = BODY_TEXT
        return p

    def add_bullet(lead, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run_lead = p.add_run(lead + " ")
        run_lead.font.bold = True
        run_lead.font.color.rgb = SECONDARY
        run_text = p.add_run(text)
        run_text.font.color.rgb = BODY_TEXT
        return p

    # --- DOCUMENT GENERATION ---
    add_doc_header()

    # SECTION 1
    add_h1("1. Executive Summary & Business Architecture")
    add_p("MediUnity operates a multi-sided digital healthcare marketplace model that monetizes transactions, platform software access, advertising placements, and enterprise healthcare services. Designed for global scalability, the platform generates recurring and transactional revenue from five key stakeholder groups: Patients, Doctors, Hospitals, Diagnostic Laboratories, and Retail Pharmacies.")
    add_p("The platform's financial engine is built upon automated payment splits via Stripe, tiered SaaS subscriptions, percentage commissions on diagnostic and pharmacy orders, B2B hospital advertising placements, and corporate wellness contracts. This multi-stream revenue approach minimizes single-source dependency and ensures predictable, scalable cash flow.")

    add_callout_box(doc, [
        "Hybrid Monetization Model: Combines Transactional Take Rates (15% per consultation), SaaS Subscriptions ($29–$99/month for doctors), Partner Marketplace Commissions (5–12% on orders), B2B Advertising, and Corporate Enterprise Packages.",
        "Automated Financial Clearing: Powered by backend algorithms (in earningsController.js and paymentRouter.js) that automatically calculate platform commission, record doctor payouts, and maintain transparent audit trails."
    ], title="CORE MONETIZATION HIGHLIGHT")

    # SECTION 2
    add_h1("2. Breakdown of Primary Revenue Streams")
    add_p("MediUnity implements seven distinct, complementary revenue streams integrated directly into the software architecture:")

    add_h2("2.1 Stream 1: Telehealth Consultation Commission Split (Transactional Take Rate)")
    add_bullet("Standard Commission Rate (15%):", "By default, the platform deducts a 15% commission on every paid doctor appointment processed through Stripe. For instance, on a $50 consultation, MediUnity retains $7.50 as platform revenue, while $42.50 is credited to the doctor's payout balance.")
    add_bullet("Dynamic Subscription Override (10%):", "Subscribed premium doctors receive a reduced platform commission rate of 10% as an incentive for maintaining active paid SaaS plans, encouraging long-term platform loyalty.")
    add_bullet("Backend Automated Splitting:", "Handled in real time by earningsController.js during Stripe webhook payment verification, instantly populating the Payment and Payout database models.")

    add_h2("2.2 Stream 2: Clinician SaaS Subscription Tiers (Recurring SaaS Revenue)")
    add_bullet("Basic Clinician Plan (Free):", "Standard marketplace listing, 15% commission per appointment, basic calendar scheduling.")
    add_bullet("Doctor Pro Plan ($29 / month):", "Reduced 10% commission rate, priority placement in doctor search filters, advanced patient vitals analytics, branded digital prescription header.")
    add_bullet("Doctor Elite / Practice Plan ($79 / month):", "Reduced 8% commission rate, featured badge on marketplace home, unlimited multi-clinic slot scheduling, integrated Socket.IO telehealth video room priority, SMS notification credits.")

    add_h2("2.3 Stream 3: Retail Pharmacy E-Commerce & Order Convenience Fees")
    add_bullet("Partner Pharmacy Marketplace Commission (8%–10%):", "MediUnity charges registered retail pharmacies a percentage commission on all prescription and OTC medication orders fulfilled through the platform.")
    add_bullet("Prescription Processing & Handling Fee ($1.50 per order):", "A nominal convenience fee charged to patients for digital prescription OCR parsing (Tesseract.js) and order routing.")
    add_bullet("Home Delivery Convenience Margin ($2.00–$4.00):", "Added margin fee shared between MediUnity and logistics delivery riders.")

    add_h2("2.4 Stream 4: Diagnostic Laboratory Test Commissions & Home Sample Fees")
    add_bullet("Diagnostic Test Booking Take Rate (10%–12%):", "Diagnostic centers pay a percentage fee for every blood test, MRI, CT scan, or ultrasound booked through MediUnity.")
    add_bullet("Home Specimen Collection Fee ($3.00–$5.00):", "Convenience fee split for arranging certified phlebotomist home visits for blood sample collection.")

    add_h2("2.5 Stream 5: Hospital B2B Sponsored Advertising & Featured Listings")
    add_bullet("Hospital Banner Advertisements (HospitalAd Model):", "Hospitals and specialized clinics pay monthly sponsorship fees to feature promotional banners on the patient portal homepage and specialty search pages.")
    add_bullet("Featured Search Position Bidding:", "Hospitals bid for top placement in regional searches for specialized departments (e.g., Cardiac Emergency, ICU Bed availability).")
    add_bullet("Inpatient Bed Booking & Transfer Fee ($15–$50 per referral):", "Success-based referral fee charged to hospitals upon patient admission through MediUnity's emergency inquiry portal.")

    add_h2("2.6 Stream 6: Home Healthcare & Emergency Service Commissions")
    add_bullet("Home Care Nursing & Physiotherapy Fee (10%):", "Commission on home nursing, post-surgical care, and elderly care service bookings.")
    add_bullet("Emergency Ambulance Dispatch Fee ($5 per booking):", "Platform routing fee for connecting patients with verified emergency transport services.")

    add_h2("2.7 Stream 7: B2B Corporate Wellness & Insurance Plans")
    add_bullet("Corporate Employee Wellness Packages ($5–$12 / employee / month):", "Employer-sponsored plans providing staff with unlimited telehealth consultations, discounted lab tests, and home medicine delivery.")
    add_bullet("Insurance Partner API Integrations:", "B2B licensing fees paid by health insurance providers to integrate MediUnity's EHR and instant claim verification APIs.")

    # SECTION 3
    add_h1("3. Feature Matrix of Revenue Streams")
    add_p("The table below summarizes the operational mechanism and target margins for each monetization channel:")

    table = doc.add_table(rows=8, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = ["Revenue Channel", "Target Customer / Partner", "Pricing Mechanism", "Platform Margin / Rate"]
    widths = [Inches(1.6), Inches(1.5), Inches(1.8), Inches(1.6)]

    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr_cells[i].width = widths[i]
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(hdr_cells[i], "0F766E")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=120, right=120)

    rev_data = [
        ("Telehealth Commissions", "Patients & Doctors", "Per-appointment Stripe split", "15% standard (10% premium)"),
        ("Clinician Subscriptions", "Medical Doctors", "Monthly / Annual SaaS Tiers", "$29 – $79 / doctor / month"),
        ("Pharmacy E-Commerce", "Retail Pharmacies", "Order take-rate + handling fee", "8% – 10% per order + $1.50 fee"),
        ("Diagnostic Lab Bookings", "Diagnostic Centers", "Lab test booking commission", "10% – 12% per lab test"),
        ("Hospital Advertising", "Hospitals & Clinics", "HospitalAd monthly sponsored slots", "$150 – $500 / hospital / month"),
        ("Home Health & Ambulance", "Care Providers & Drivers", "Booking commission & dispatch fee", "10% on services + $5 ambulance"),
        ("Corporate Enterprise", "Employers & Insurers", "Per-employee monthly SaaS fee", "$5 – $12 / employee / month")
    ]

    for row_idx, data in enumerate(rev_data, start=1):
        row_cells = table.rows[row_idx].cells
        bg_hex = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            row_cells[col_idx].width = widths[col_idx]
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.font.size = Pt(9.5)
            if col_idx == 0:
                run.font.bold = True
                run.font.color.rgb = SECONDARY
            elif col_idx == 3:
                run.font.bold = True
                run.font.color.rgb = PRIMARY
            else:
                run.font.color.rgb = BODY_TEXT
            set_cell_background(row_cells[col_idx], bg_hex)
            set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=120, right=120)

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(12)

    # SECTION 4
    add_h1("4. Pro Forma Financial Revenue Projections")
    add_p("To demonstrate commercial viability, the table below projects estimated monthly and annual revenue based on moderate platform scale (500 Active Doctors, 10,000 Monthly Appointments, 2,500 Pharmacy Orders, 1,500 Lab Tests, and 20 Hospital Sponsors):")

    table_proj = doc.add_table(rows=7, cols=4)
    table_proj.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_proj.autofit = False

    proj_headers = ["Revenue Category", "Monthly Volume / Units", "Average Unit Price / Take", "Est. Monthly Revenue"]
    proj_widths = [Inches(1.8), Inches(1.6), Inches(1.6), Inches(1.5)]

    p_hdr_cells = table_proj.rows[0].cells
    for i, text in enumerate(proj_headers):
        p_hdr_cells[i].width = proj_widths[i]
        p = p_hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(p_hdr_cells[i], "0F766E")
        set_cell_margins(p_hdr_cells[i], top=120, bottom=120, left=120, right=120)

    proj_data = [
        ("Doctor Consultations", "10,000 appointments", "$40 avg fee ($6.00 take @ 15%)", "$60,000 / month"),
        ("Doctor Pro Subscriptions", "200 subscribed doctors", "$49 / month average plan", "$9,800 / month"),
        ("Pharmacy Orders", "2,500 medicine orders", "$30 avg order ($3.00 take @ 10%)", "$7,500 / month"),
        ("Diagnostic Lab Tests", "1,500 lab bookings", "$50 avg test ($6.00 take @ 12%)", "$9,000 / month"),
        ("Hospital Ads & B2B Leads", "20 hospital sponsors", "$300 / hospital ad slot", "$6,000 / month"),
        ("TOTAL ESTIMATED REVENUE", "Scale: 500 Doctors / 10k Patients", "Combined Revenue Engine", "$92,300 / month ($1.1M / yr)")
    ]

    for row_idx, data in enumerate(proj_data, start=1):
        row_cells = table_proj.rows[row_idx].cells
        bg_hex = "F8FAFC" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            row_cells[col_idx].width = proj_widths[col_idx]
            p = row_cells[col_idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            run.font.size = Pt(9.5)
            if row_idx == 6:
                run.font.bold = True
                run.font.color.rgb = PRIMARY if col_idx == 3 else SECONDARY
            elif col_idx == 0:
                run.font.bold = True
                run.font.color.rgb = SECONDARY
            else:
                run.font.color.rgb = BODY_TEXT
            set_cell_background(row_cells[col_idx], "E0F2FE" if row_idx == 6 else bg_hex)
            set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=120, right=120)

    p_space2 = doc.add_paragraph()
    p_space2.paragraph_format.space_before = Pt(12)

    # SECTION 5
    add_h1("5. Competitive Strengths & Financial Growth Strategy")
    add_bullet("Asset-Light Enterprise Model:", "MediUnity does not own physical clinics, laboratories, or warehouses. It operates as a high-margin digital orchestrator, resulting in minimal operational overhead.")
    add_bullet("Global Multi-Currency Scaling:", "Integrated Stripe payment pipelines allow instant multi-currency processing (USD, EUR, GBP, BDT, INR), permitting seamless expansion into new international markets.")
    add_bullet("High Customer Retention:", "Centralized Electronic Health Records (EHR) and continuous vitals graphing create high switching costs for patients and doctors, driving recurring long-term revenue.")

    # SECTION 6
    add_h1("6. Conclusion")
    add_p("MediUnity's business monetization model presents a highly attractive, scalable, and diversified financial structure. By combining transaction commissions, SaaS subscriptions, pharmacy margins, lab test take rates, and hospital advertising, MediUnity achieves sustainable revenue growth while expanding affordable, high-quality digital healthcare worldwide.")

    # Save Document
    doc.save(output_path)
    print(f"Successfully generated monetization document at: {output_path}")

if __name__ == "__main__":
    downloads_dir = os.path.join(os.path.expanduser("~"), "Downloads")
    os.makedirs(downloads_dir, exist_ok=True)
    target_file = os.path.join(downloads_dir, "MediUnity_Monetization_and_Revenue_Model.docx")
    build_monetization_document(target_file)
